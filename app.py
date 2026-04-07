# app.py - Cleaned version with AI functions imported from models.py, workspace from workspace.py, and docs from mydocs.py

import os
import json
import re
import uuid
import threading
import time
import asyncio
import aiohttp
from flask import Flask, render_template, request, jsonify, send_file, Response
from flask_cors import CORS
from werkzeug.utils import secure_filename
import requests
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Inches
import base64
from binary_processor import BinaryProcessor
from docs import DocumentProcessor
import io

# Import AI functions from models.py
from models import (
    query_ai_with_fallback,
    generate_chat_title,
    is_code_generation_request,
    execute_python_code,
    search_web,
    extract_web_content,
    analyze_image_with_ai,
    call_pollinations_ai
)

# Import workspace functions from workspace.py
from workspace import register_workspace_routes

# Import free image generator
from image import FreeImageGenerator

# Import document creation utilities from mydocs.py
from mydocs import DocumentCreator

from vision import get_vision_model

# Import terminal blueprint
from terminal import create_terminal_blueprint

app = Flask(__name__)
app.secret_key = os.urandom(24)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # Increase to 50MB
CORS(app)

# Register terminal routes
app.register_blueprint(create_terminal_blueprint(app))

# Ensure the generated_images directory exists
import os
os.makedirs('generated_images', exist_ok=True)

binary_processor = BinaryProcessor()

# Initialize free image generator (NO API KEYS NEEDED!)
image_generator = FreeImageGenerator(output_dir="generated_images")

CONVERSATIONS_FILE = 'conversations.json'

# ============= WAKEUP PING SYSTEM FOR RENDER =============
# This prevents the app from sleeping after 15 minutes of inactivity

# Configuration
WAKEUP_ENABLED = os.environ.get("WAKEUP_ENABLED", "true").lower() == "true"
PING_INTERVAL = int(os.environ.get("PING_INTERVAL", "300"))  # 5 minutes default (300 seconds)
PING_URLS = []  # Will be populated after server starts

# List of free monitoring services (fallback options)
MONITORING_SERVICES = [
    "https://uptimerobot.com",  # External monitoring (requires account)
    "https://betteruptime.com",  # External monitoring (requires account)
    "https://statuscake.com",    # External monitoring (requires account)
]

def get_local_url():
    """Get the local URL of the app"""
    port = os.environ.get("PORT", "5001")
    # For Render, the app is accessible via the assigned URL
    # We'll use the RENDER_EXTERNAL_URL if available
    render_url = os.environ.get("RENDER_EXTERNAL_URL")
    if render_url:
        return render_url
    return f"http://localhost:{port}"

def wakeup_ping():
    """Send a ping to keep the app alive"""
    urls_to_ping = PING_URLS.copy()
    
    # Add local URL if not already present
    local_url = get_local_url()
    if local_url not in urls_to_ping:
        urls_to_ping.append(local_url)
    
    # Also ping the health endpoint
    health_endpoint = f"{local_url}/api/health"
    if health_endpoint not in urls_to_ping:
        urls_to_ping.append(health_endpoint)
    
    success_count = 0
    failed_count = 0
    
    for url in urls_to_ping:
        try:
            # Use a short timeout to avoid hanging
            response = requests.get(url, timeout=5)
            if response.status_code < 500:
                success_count += 1
                print(f"✅ Wakeup ping successful: {url} - Status: {response.status_code}")
            else:
                failed_count += 1
                print(f"⚠️ Wakeup ping warning: {url} - Status: {response.status_code}")
        except requests.RequestException as e:
            failed_count += 1
            print(f"❌ Wakeup ping failed: {url} - Error: {str(e)}")
    
    return success_count, failed_count

def async_wakeup_ping():
    """Async version of wakeup ping using aiohttp (fallback)"""
    async def _ping():
        urls_to_ping = PING_URLS.copy()
        local_url = get_local_url()
        if local_url not in urls_to_ping:
            urls_to_ping.append(local_url)
        
        health_endpoint = f"{local_url}/api/health"
        if health_endpoint not in urls_to_ping:
            urls_to_ping.append(health_endpoint)
        
        async with aiohttp.ClientSession() as session:
            for url in urls_to_ping:
                try:
                    async with session.get(url, timeout=aiohttp.ClientTimeout(total=5)) as response:
                        if response.status < 500:
                            print(f"✅ Async wakeup ping successful: {url}")
                        else:
                            print(f"⚠️ Async wakeup ping warning: {url}")
                except Exception as e:
                    print(f"❌ Async wakeup ping failed: {url} - Error: {str(e)}")
    
    try:
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        loop.run_until_complete(_ping())
        loop.close()
    except Exception as e:
        print(f"⚠️ Async wakeup ping error: {e}")

def start_wakeup_scheduler():
    """Start the background scheduler for wakeup pings"""
    if not WAKEUP_ENABLED:
        print("⚠️ Wakeup pings are disabled. Set WAKEUP_ENABLED=true to enable.")
        return
    
    def run_scheduler():
        print(f"🚀 Wakeup ping scheduler started. Pinging every {PING_INTERVAL} seconds.")
        print(f"📍 Local URL: {get_local_url()}")
        
        # Do an initial ping immediately
        wakeup_ping()
        
        while True:
            time.sleep(PING_INTERVAL)
            try:
                # Try primary method first
                success, failed = wakeup_ping()
                
                # If primary method had failures, try async fallback
                if failed > 0:
                    print("⚠️ Primary ping had failures, trying async fallback...")
                    async_wakeup_ping()
                
                # Log status
                total_attempts = success + failed
                if total_attempts > 0:
                    success_rate = (success / total_attempts) * 100
                    print(f"📊 Ping stats: {success}/{total_attempts} successful ({success_rate:.1f}%)")
                
            except Exception as e:
                print(f"❌ Wakeup scheduler error: {e}")
                # Try fallback on error
                try:
                    async_wakeup_ping()
                except:
                    pass
    
    # Start the scheduler in a background daemon thread
    scheduler_thread = threading.Thread(target=run_scheduler, daemon=True)
    scheduler_thread.start()
    return scheduler_thread

# ============= ADDITIONAL KEEP-ALIVE MECHANISMS =============

# Mechanism 2: Self-ping using internal requests every 4 minutes (if no external traffic)
def internal_self_ping():
    """Internal self-ping using Flask test client"""
    try:
        with app.test_client() as client:
            response = client.get('/api/health')
            if response.status_code < 500:
                print(f"✅ Internal self-ping successful: {response.status_code}")
            else:
                print(f"⚠️ Internal self-ping warning: {response.status_code}")
    except Exception as e:
        print(f"❌ Internal self-ping failed: {e}")

def start_internal_ping_scheduler():
    """Start internal ping scheduler as backup"""
    def run_internal_pings():
        # Ping every 4 minutes (240 seconds) as fallback
        while True:
            time.sleep(240)
            try:
                internal_self_ping()
            except Exception as e:
                print(f"⚠️ Internal ping error: {e}")
    
    thread = threading.Thread(target=run_internal_pings, daemon=True)
    thread.start()
    return thread

# Mechanism 3: Keep connections alive using a simple counter
_keepalive_counter = 0

@app.before_request
def track_request():
    """Track requests to keep the app active"""
    global _keepalive_counter
    _keepalive_counter += 1
    if _keepalive_counter % 100 == 0:
        print(f"📊 Request counter: {_keepalive_counter} total requests")

# Mechanism 4: Health check endpoint that also triggers activity
@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint - also keeps the app alive"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'uptime': time.time() - app.config.get('start_time', time.time()),
        'request_count': _keepalive_counter,
        'wakeup_enabled': WAKEUP_ENABLED,
        'ping_interval': PING_INTERVAL
    })

# Mechanism 5: Simple root endpoint with timestamp
@app.route('/api/ping', methods=['GET'])
def ping_endpoint():
    """Simple ping endpoint for external monitoring services"""
    return jsonify({
        'pong': True,
        'timestamp': datetime.now().isoformat(),
        'server_time': time.time()
    })

# Mechanism 6: Background thread that keeps the database connection alive (if applicable)
def keep_database_alive():
    """Keep any database connections alive (if you add a database later)"""
    # This is a placeholder for future database keep-alive
    pass

def start_db_keepalive():
    """Start database keep-alive thread"""
    def run_db_keepalive():
        while True:
            time.sleep(300)  # Every 5 minutes
            try:
                keep_database_alive()
            except Exception as e:
                print(f"⚠️ DB keep-alive error: {e}")
    
    thread = threading.Thread(target=run_db_keepalive, daemon=True)
    thread.start()
    return thread

# ============= INITIALIZE ALL KEEP-ALIVE MECHANISMS =============

# Store start time for uptime tracking
app.config['start_time'] = time.time()

# Start all keep-alive mechanisms
wakeup_thread = start_wakeup_scheduler()
internal_ping_thread = start_internal_ping_scheduler()
db_keepalive_thread = start_db_keepalive()

print("\n" + "="*60)
print("🔄 KEEP-ALIVE SYSTEM INITIALIZED")
print("="*60)
print(f"   Wakeup Enabled: {WAKEUP_ENABLED}")
print(f"   Ping Interval: {PING_INTERVAL} seconds")
print(f"   Internal Pings: Enabled (every 240 seconds)")
print(f"   Health Endpoint: /api/health")
print(f"   Ping Endpoint: /api/ping")
print("="*60 + "\n")

# ============= END WAKEUP PING SYSTEM =============

# Allowed file extensions for documents (expanded for binary files)
ALLOWED_EXTENSIONS = {
    # Text formats
    'txt', 'md', 'py', 'js', 'html', 'css', 'json', 'xml', 'csv',
    # Documents
    'pdf', 'doc', 'docx', 'xls', 'xlsx', 'ppt', 'pptx',
    # Media
    'jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff', 'webp',
    'mp3', 'wav', 'ogg', 'flac', 'm4a',
    'mp4', 'avi', 'mov', 'mkv', 'webm',
    # Archives
    'zip', 'rar', '7z', 'tar', 'gz', 'bz2',
    # Executables (limited)
    'exe', 'dll', 'so', 'dylib',
    # Databases
    'db', 'sqlite', 'sqlite3', 'sql',
    # Fonts
    'ttf', 'otf', 'woff', 'woff2',
    # Code
    'java', 'c', 'cpp', 'h', 'rb', 'php', 'go', 'rs', 'swift', 'kt'
}

def sanitize_filename(text):
    """Sanitize text to be safe for use in filenames"""
    if not text:
        return "New Chat"
    # Remove control characters
    text = ''.join(char for char in text if char.isprintable() and char not in '\n\r\t')
    # Replace problematic characters
    replacements = {
        '/': '-',
        '\\': '-',
        ':': '-',
        '*': '-',
        '?': '-',
        '"': "'",
        '<': '-',
        '>': '-',
        '|': '-'
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    # Trim whitespace and limit length
    text = text.strip()[:100]
    return text if text else "New Chat"

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def load_conversations():
    try:
        if os.path.exists(CONVERSATIONS_FILE):
            with open(CONVERSATIONS_FILE, 'r', encoding='utf-8') as f:
                conversations = json.load(f)
                for conv_id, conv_data in conversations.items():
                    if 'versions' not in conv_data:
                        conv_data['versions'] = {}
                    if 'current_version_index' not in conv_data:
                        conv_data['current_version_index'] = {}
                    if 'branch_root' not in conv_data:
                        conv_data['branch_root'] = None
                return conversations
    except Exception as e:
        print(f"Error loading conversations: {e}")
    return {}

def save_conversations(conversations):
    try:
        with open(CONVERSATIONS_FILE, 'w', encoding='utf-8') as f:
            json.dump(conversations, f, indent=2, ensure_ascii=False)
    except Exception as e:
        print(f"Error saving conversations: {e}")

conversations = load_conversations()

def extract_text_from_file(file_content, filename):
    """Extract text from uploaded files using enhanced binary processor"""
    try:
        # Use the enhanced binary processor
        processed_output = binary_processor.process_file(file_content, filename)
        
        # For AI consumption, limit to reasonable size (50KB should be enough for AI)
        if len(processed_output) > 50000:
            processed_output = processed_output[:50000] + "\n\n[Content truncated due to size]"
        
        return processed_output
        
    except Exception as e:
        print(f"Error in enhanced extraction: {e}")
        # Fallback to original method
        return f"[Error extracting from {filename}: {str(e)}]"

def export_to_word(conversation_data):
    """Export conversation to Word document"""
    doc = Document()
    doc.add_heading(f'HenAi Chat: {conversation_data["title"]}', 0)
    doc.add_paragraph(f'Exported on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph('-' * 50)

    for msg in conversation_data['messages']:
        if msg['role'] == 'user':
            doc.add_heading('You:', level=2)
        else:
            doc.add_heading('HenAi:', level=2)

        # Handle code blocks in content
        content = msg['content']
        # Extract code blocks
        code_blocks = re.findall(r'```(\w+)?\n([\s\S]*?)```', content)
        if code_blocks:
            # Split content by code blocks
            parts = re.split(r'```\w*\n[\s\S]*?```', content)
            for i, part in enumerate(parts):
                if part.strip():
                    doc.add_paragraph(part.strip())
                if i < len(code_blocks):
                    lang, code = code_blocks[i]
                    doc.add_paragraph(f'[{lang.upper()} Code]')
                    doc.add_paragraph(code.strip())
        else:
            doc.add_paragraph(content)
        doc.add_paragraph('')

    return doc

def is_image_generation_request(message):
    """
    Detect if a user message is requesting image generation.
    Returns True if the message contains keywords indicating image generation.
    """
    if not message:
        return False
    
    message_lower = message.lower()
    
    # First, check if this is an analysis request - these should NOT trigger image generation
    analysis_keywords = [
        'analyze', 'analyse', 'what is', 'what\'s', 'tell me about', 
        'describe', 'explain', 'read this', 'look at', 'examine',
        'extract text', 'ocr', 'recognize', 'identify'
    ]
    
    for keyword in analysis_keywords:
        if keyword in message_lower:
            return False
    
    # Image generation keywords
    image_keywords = [
        'generate image', 'create image', 'make image', 'draw image',
        'generate picture', 'create picture', 'make picture', 'draw picture',
        'generate photo', 'create photo', 'make photo',
        'generate art', 'create art', 'make art',
        'generate illustration', 'create illustration',
        'ai image', 'ai art', 'ai generate',
        'stable diffusion', 'dall-e', 'midjourney',
        'flux', 'sdxl', 'sd1.5',
        'image of', 'picture of', 'photo of',
        'draw me', 'generate me', 'create me',
        'make an image', 'make a picture', 'generate an image'
    ]
    
    # Check if any keyword is in the message
    for keyword in image_keywords:
        if keyword in message_lower:
            return True
    
    # Additional check for short descriptive phrases that might be image requests
    # BUT only if the message is very short (less than 5 words) and doesn't contain analysis terms
    words = message_lower.split()
    if len(words) <= 5 and any(word in ['image', 'picture', 'photo', 'art'] for word in words):
        # Also check that it's not asking about an existing image
        if not any(word in message_lower for word in ['this', 'the', 'that', 'attached']):
            return True
    
    return False

def extract_image_prompt(message):
    """
    Extract the actual image prompt from the message by removing command prefixes.
    """
    message_lower = message.lower()
    
    # Remove common prefixes
    prefixes = [
        'generate image of', 'create image of', 'make image of', 'draw image of',
        'generate picture of', 'create picture of', 'make picture of', 'draw picture of',
        'generate photo of', 'create photo of', 'make photo of',
        'generate art of', 'create art of', 'make art of',
        'ai image of', 'ai art of',
        'image of', 'picture of', 'photo of',
        'draw me', 'generate me', 'create me',
        'generate an image of', 'create an image of', 'make an image of',
        'generate a picture of', 'create a picture of', 'make a picture of'
    ]
    
    prompt = message.strip()
    for prefix in prefixes:
        if message_lower.startswith(prefix):
            prompt = message[len(prefix):].strip()
            break
    
    # Remove any trailing punctuation and clean up
    prompt = prompt.strip('.,!?;:')
    
    # If prompt is empty or too short, use the original message
    if len(prompt) < 3:
        prompt = message.strip()
    
    return prompt

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/chat', methods=['POST'])
def chat():
    data = request.json
    message = data.get('message', '')
    conversation_id = data.get('conversation_id')
    regenerate = data.get('regenerate', False)
    regenerate_from = data.get('regenerate_from')
    attached_files = data.get('attached_files', [])

    if not conversation_id or conversation_id not in conversations:
        conversation_id = str(uuid.uuid4())
        conversations[conversation_id] = {
            'id': conversation_id,
            'messages': [],
            'title': 'New Chat',
            'created_at': datetime.now().isoformat(),
            'last_updated': datetime.now().isoformat(),
            'versions': {},
            'current_version_index': {},
            'branch_root': None
        }

    conv = conversations[conversation_id]

    if 'versions' not in conv:
        conv['versions'] = {}
    if 'current_version_index' not in conv:
        conv['current_version_index'] = {}
    if 'branch_root' not in conv:
        conv['branch_root'] = None

    # ===== CHECK FOR PENDING ATTACHMENTS FROM ARCHIVE =====
    if 'pending_attachments' in conv and conv['pending_attachments']:
        # If user is sending a message, attach pending files to this message
        if message:
            if not attached_files:
                attached_files = []
        
        # Add all pending attachments
        for pending_file in conv['pending_attachments']:
            attached_files.append({
                'name': pending_file['name'],
                'content': pending_file['content']
            })
        
        # Clear pending attachments after they're used
        conv['pending_attachments'] = []
        
        # Save the conversation
        save_conversations(conversations)
    # ===== END PENDING ATTACHMENTS CHECK =====

    # Handle regeneration
    if regenerate and regenerate_from is not None:
        msg_key = str(regenerate_from)

        if msg_key not in conv['versions']:
            conv['versions'][msg_key] = []
        
        # Initialize version_branches if not exists
        if 'version_branches' not in conv:
            conv['version_branches'] = {}

        # Save the current response as a version BEFORE truncating
        if regenerate_from < len(conv['messages']):
            current_response = conv['messages'][regenerate_from]['content']
            # Get current version index before adding new version
            current_version_idx = conv['current_version_index'].get(msg_key, 0)
            
            # Only save if it's different from existing versions
            if current_response not in conv['versions'][msg_key]:
                conv['versions'][msg_key].append(current_response)
                # Set index to point to the version we just saved (the original)
                conv['current_version_index'][msg_key] = len(conv['versions'][msg_key]) - 1
                current_version_idx = len(conv['versions'][msg_key]) - 1
            
            # CRITICAL FIX: Save the current branch's subsequent messages before truncating
            # This preserves the conversation history for the current version
            subsequent_messages = conv['messages'][regenerate_from + 1:]
            if subsequent_messages:
                version_branch_key = f"{msg_key}_v{current_version_idx}"
                conv['version_branches'][version_branch_key] = subsequent_messages

        # Truncate messages at the user message level (regenerate_from - 1)
        # This creates a new branch point
        if regenerate_from > 0 and regenerate_from - 1 < len(conv['messages']):
            user_message_content = conv['messages'][regenerate_from - 1]['content']
            # Truncate to keep only up to the user message (not including the old assistant response)
            conv['messages'] = conv['messages'][:regenerate_from]
            conv['branch_root'] = regenerate_from - 1
            
        # Set the new version index for the regenerated response
        new_version_idx = len(conv['versions'].get(msg_key, []))  # Will be the index of the new response
        conv['current_version_index'][msg_key] = new_version_idx
            
        # Clear any version indices for messages that no longer exist after truncation
        indices_to_remove = []
        for idx_key in list(conv.get('current_version_index', {}).keys()):
            try:
                if int(idx_key) >= regenerate_from:
                    indices_to_remove.append(idx_key)
            except ValueError:
                continue
        
        for idx_key in indices_to_remove:
            del conv['current_version_index'][idx_key]

    # Handle file attachments - send content to AI, show only names to user
    full_message = message
    ai_message = message
    display_message = message
    
    if attached_files:
        # Build content for AI (includes full file content)
        file_contexts = []
        file_names = []
        for file_info in attached_files:
            file_names.append(file_info['name'])
            file_contexts.append(f"\n\n--- BEGIN FILE: {file_info['name']} ---\n{file_info['content']}\n--- END FILE: {file_info['name']} ---\n")
        
        # For AI: include full file content
        ai_message = message + ''.join(file_contexts)
        # For display: only show file names (no content)
        file_list = ", ".join(file_names)
        display_message = message + (f"\n\n📎 **Attached files:** {file_list}" if message else f"📎 **Attached files:** {file_list}")
        
        full_message = display_message
    else:
        ai_message = message
        full_message = message

    response_data = {
        'response': '',
        'code_execution': None,
        'conversation_id': conversation_id,
        'title': conv['title'],
        'versions': conv.get('versions', {}),
        'current_version_index': conv.get('current_version_index', {})
    }

    # Handle commands
    if full_message.lower().startswith('/search'):
        query = full_message[7:].strip()
        if not query:
            response_data['response'] = "Please provide a search query. Example: `/search artificial intelligence`"
        else:
            response_data['response'] = search_web(query)

    elif full_message.lower().startswith('/extract'):
        url = full_message[8:].strip()
        if not url:
            response_data['response'] = "Please provide a URL. Example: `/extract https://example.com`"
        else:
            response_data['response'] = extract_web_content(url)

    elif full_message.lower().startswith('/code'):
        code = full_message[5:].strip()
        if not code:
            response_data['response'] = "Please provide Python code. Example: `/code print('Hello World!')`"
        else:
            exec_result = execute_python_code(code)
            response_data['code_execution'] = exec_result
            if exec_result['success']:
                response_data['response'] = f"✅ **Code executed successfully!**\n\n```\n{exec_result['output']}\n```"
            else:
                response_data['response'] = f"❌ **Code execution failed!**\n\n```\n{exec_result['error']}\n```"

    # NEW: Handle image generation with FreeImageGenerator
    elif full_message.lower().startswith('/generate') or full_message.lower().startswith('/image'):
        # Extract the image prompt
        if full_message.lower().startswith('/generate'):
            image_prompt = full_message[9:].strip()
        else:
            image_prompt = full_message[6:].strip()
        
        if not image_prompt:
            response_data['response'] = "Please provide an image description. Example: `/generate a beautiful sunset over mountains`"
        else:
            try:
                # Generate image using free methods (no API keys)
                import os
                # datetime is already imported at the top, no need to import again
                
                # Generate a unique filename
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"ai_gen_{timestamp}.png"
                
                # Try generating with available methods (prefer Hugging Face API first)
                try:
                    # Try Hugging Face API first (faster, no model download)
                    image_path = image_generator.generate_huggingface(image_prompt, output_name=filename)
                except Exception as e:
                    print(f"Hugging Face generation failed: {e}, falling back to local generation")
                    # Fallback to local generation if available
                    try:
                        image_path = image_generator.generate_local_sd(image_prompt, output_name=filename)
                    except Exception as e2:
                        print(f"Local generation also failed: {e2}")
                        raise e
                
                # Get the URL for the generated image
                # Serve the image from the generated_images directory
                import os
                image_filename = os.path.basename(image_path)
                
                # Create a response with image data and controls
                image_html = f'''
<div style="text-align: center; margin: 15px 0;">
    <img src="/api/generated_image/{image_filename}" 
         alt="Generated: {image_prompt}" 
         style="max-width: 100%; max-height: 400px; border-radius: 12px; cursor: pointer; box-shadow: 0 4px 12px rgba(0,0,0,0.3);"
         onclick="window.open('/api/generated_image/{image_filename}', '_blank')">
    <div style="margin-top: 8px; font-size: 0.75rem; color: var(--text-muted);">
        <i class="fas fa-expand"></i> Click to view full size
    </div>
</div>
'''
                response_data['response'] = f"🎨 **Generated Image: \"{image_prompt}\"**\n\n{image_html}"
                response_data['is_image_generation'] = True
                response_data['image_path'] = image_path
                response_data['image_prompt'] = image_prompt
                
            except Exception as e:
                print(f"Image generation error: {e}")
                response_data['response'] = f"❌ **Image generation failed!**\n\nError: {str(e)}\n\nPlease try a different prompt or use the /image command with a more detailed description."

    elif full_message.lower().startswith('/help'):
        response_data['response'] = """**📚 HenAi Commands & Features**

**Commands:**
• `/search <query>` - Search the web
• `/extract <url>` - Extract content from a URL
• `/code <python>` - Execute Python code
• `/generate <description>` or `/image <description>` - Generate AI images (free!)
• `/help` - Show this help

**Message Actions:**
• 📋 **Copy** - Click copy on any message
• ✏️ **Edit** - Click edit on your messages
• 🔄 **Regenerate** - Get new responses
• ↔️ **Toggle Versions** - Switch between chat versions (creates branches)

**Features:**
• 📎 Multiple file attachments (text, code, documents)
• 💾 Auto-save conversations
• 🏷️ Auto-titled chats
• ✏️ Rename chats
• 📥 Export chats to Word
• 🔄 Version history with branching
• 🎨 Free AI Image Generation (no API keys needed!)"""
    
    # NEW: Auto-detect image generation from natural language
    # BUT skip if there are file attachments (which should be analyzed, not used for generation)
    elif is_image_generation_request(full_message) and not attached_files:
        # This is a natural language image generation request (no attachments)
        image_prompt = extract_image_prompt(full_message)
        
        try:
            # Generate image using free methods (no API keys)
            import os
            
            # Generate a unique filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"ai_gen_{timestamp}.png"
            
            # Try generating with available methods
            try:
                # Try Hugging Face API first (faster, no model download)
                image_path = image_generator.generate_huggingface(image_prompt, output_name=filename)
            except Exception as e:
                print(f"Hugging Face generation failed: {e}, falling back to local generation")
                # Fallback to local generation if available
                try:
                    image_path = image_generator.generate_local_sd(image_prompt, output_name=filename)
                except Exception as e2:
                    print(f"Local generation also failed: {e2}")
                    raise e
            
            # Get the URL for the generated image
            import os
            image_filename = os.path.basename(image_path)
            
            # Create a response with image data and controls
            image_html = f'''
<div style="text-align: center; margin: 15px 0;">
    <img src="/api/generated_image/{image_filename}" 
         alt="Generated: {image_prompt}" 
         style="max-width: 100%; max-height: 400px; border-radius: 12px; cursor: pointer; box-shadow: 0 4px 12px rgba(0,0,0,0.3);"
         onclick="window.open('/api/generated_image/{image_filename}', '_blank')">
    <div style="margin-top: 8px; font-size: 0.75rem; color: var(--text-muted);">
        <i class="fas fa-expand"></i> Click to view full size
    </div>
</div>
'''
            response_data['response'] = f"🎨 **Here's an image I generated for you:**\n\n{image_html}"
            response_data['is_image_generation'] = True
            response_data['image_path'] = image_path
            response_data['image_prompt'] = image_prompt
            
        except Exception as e:
            print(f"Image generation error: {e}")
            response_data['response'] = f"❌ **Image generation failed!**\n\nError: {str(e)}\n\nPlease try a different prompt or use the `/generate` command."

    else:
        # Build full context from all messages with proper file content
        context = []
        for msg in conv['messages']:
            # For user messages with attachments, we need to include the actual content
            # that was sent to the AI, not just the display version
            if msg['role'] == 'user' and msg.get('attachments'):
                # Check if we have the AI-ready content stored
                if 'ai_content' in msg:
                    context.append({"role": msg['role'], "content": msg['ai_content']})
                else:
                    # Reconstruct from message and attachments
                    user_content = msg['content']
                    if 'attachments' in msg and msg['attachments']:
                        file_context = ""
                        for file_info in msg['attachments']:
                            if 'content' in file_info:
                                file_context += file_info['content']  # Already formatted
                        if file_context:
                            user_content = msg['content'] + file_context
                    context.append({"role": msg['role'], "content": user_content})
            else:
                context.append({"role": msg['role'], "content": msg['content']})
        
        # Debug logging to verify context
        print(f"Context messages count: {len(context)}")
        for i, msg in enumerate(context[-5:]):  # Show last 5 messages
            print(f"  Message {i}: role={msg['role']}, content length={len(msg['content'])}")

        # Check if this is a code generation request (improved detection)
        is_code_gen = is_code_generation_request(full_message)
        
        # Log the decision for debugging
        print(f"Code generation detection for '{full_message[:50]}...': {is_code_gen}")

        # Only add the code generation instruction for clear code requests
        prompt_for_ai = ai_message if 'ai_message' in locals() else full_message
        
        if is_code_gen:
            # Add instruction only for explicit code generation requests
            prompt_for_ai += "\n\nIMPORTANT: Provide the COMPLETE, FULLY FUNCTIONAL code with at least 500 lines. Do not abbreviate or use placeholders. Include all necessary imports, functions, error handling, and comments."

        # Use the enhanced AI query with appropriate service
        # - Code generation: OpenRouter only
        # - Non-code: Pollinations.ai first, then OpenRouter
        ai_response = query_ai_with_fallback(prompt_for_ai, context, is_code_gen)
        response_data['response'] = ai_response

    # Add messages - store both display version and AI-ready version
    if not regenerate or regenerate_from is None:
        stored_content = display_message if 'display_message' in locals() else full_message
        ai_ready_content = ai_message if 'ai_message' in locals() else full_message
        
        conv['messages'].append({
            'role': 'user',
            'content': stored_content,  # Display version for UI
            'ai_content': ai_ready_content,  # Full version for AI context
            'timestamp': datetime.now().isoformat(),
            'attachments': attached_files if attached_files else None
        })
        conv['branch_root'] = None

    new_assistant_index = len(conv['messages'])
    conv['messages'].append({
        'role': 'assistant',
        'content': response_data['response'],
        'timestamp': datetime.now().isoformat()
    })

    # Update versions for regeneration
    if regenerate and regenerate_from is not None:
        msg_key = str(regenerate_from)
        if msg_key not in conv['versions']:
            conv['versions'][msg_key] = []

        if response_data['response'] not in conv['versions'][msg_key]:
            conv['versions'][msg_key].append(response_data['response'])
            conv['current_version_index'][msg_key] = len(conv['versions'][msg_key]) - 1

    if conv['title'] == 'New Chat':
        raw_title = generate_chat_title(conv['messages'])
        conv['title'] = sanitize_filename(raw_title)
        response_data['title'] = conv['title']

    conv['last_updated'] = datetime.now().isoformat()
    save_conversations(conversations)

    return jsonify(response_data)

@app.route('/api/edit_message', methods=['POST'])
def edit_message():
    data = request.json
    conversation_id = data.get('conversation_id')
    message_index = data.get('message_index')
    new_content = data.get('new_content')
    attached_files = data.get('attached_files', [])

    if conversation_id not in conversations:
        return jsonify({'error': 'Conversation not found'}), 404

    conv = conversations[conversation_id]

    if 'versions' not in conv:
        conv['versions'] = {}
    if 'current_version_index' not in conv:
        conv['current_version_index'] = {}

    messages = conv['messages']

    if message_index >= len(messages):
        return jsonify({'error': 'Message not found'}), 404

    full_content = new_content
    if attached_files:
        # Just show that files are attached, but don't include their content in the message
        file_names = [file_info['name'] for file_info in attached_files]
        file_list = ", ".join(file_names)
        full_content = new_content + f"\n\n[Attached files: {file_list}]"

    # Save current assistant response as version
    if message_index + 1 < len(messages):
        msg_key = str(message_index + 1)
        if msg_key not in conv['versions']:
            conv['versions'][msg_key] = []
        
        # Initialize version_branches if not exists
        if 'version_branches' not in conv:
            conv['version_branches'] = {}

        old_response = messages[message_index + 1]['content']
        current_version_idx = conv['current_version_index'].get(msg_key, 0)
        
        if old_response not in conv['versions'][msg_key]:
            conv['versions'][msg_key].append(old_response)
            conv['current_version_index'][msg_key] = len(conv['versions'][msg_key]) - 1
            current_version_idx = len(conv['versions'][msg_key]) - 1
        
        # CRITICAL FIX: Save the current branch's subsequent messages before truncating
        subsequent_messages = messages[message_index + 2:]
        if subsequent_messages:
            version_branch_key = f"{msg_key}_v{current_version_idx}"
            conv['version_branches'][version_branch_key] = subsequent_messages

    # Store both display and AI-ready content
    ai_ready_content = new_content
    if attached_files:
        # Build AI-ready content with file context
        file_context = ""
        for file_info in attached_files:
            if 'content' in file_info:
                file_context += file_info['content']  # Already formatted
        if file_context:
            ai_ready_content = new_content + file_context
    
    messages[message_index]['content'] = full_content  # Display version
    messages[message_index]['ai_content'] = ai_ready_content  # AI version
    messages[message_index]['attachments'] = attached_files if attached_files else None

    while len(messages) > message_index + 1:
        messages.pop()

    context = []
    for msg in messages:
        # For user messages with attachments, use ai_content if available
        if msg['role'] == 'user' and 'ai_content' in msg:
            context.append({"role": msg['role'], "content": msg['ai_content']})
        else:
            context.append({"role": msg['role'], "content": msg['content']})
    
    # Debug logging
    print(f"Edit context messages count: {len(context)}")

    is_code_gen = is_code_generation_request(full_content)
    
    # Only add the code generation instruction if it's a clear code request
    if is_code_gen:
        if not any(phrase in full_content.lower() for phrase in ['summarize', 'explain', 'what is', 'tell me about']):
            full_content += "\n\nIMPORTANT: Provide the COMPLETE, FULLY FUNCTIONAL code with at least 500 lines. Do not abbreviate or use placeholders. Include all necessary imports, functions, error handling, and comments."
    
    ai_response = query_ai_with_fallback(full_content, context, is_code_gen)

    messages.append({
        'role': 'assistant',
        'content': ai_response,
        'timestamp': datetime.now().isoformat()
    })

    msg_key = str(message_index + 1)
    if msg_key not in conv['versions']:
        conv['versions'][msg_key] = []

    if ai_response not in conv['versions'][msg_key]:
        conv['versions'][msg_key].append(ai_response)
        conv['current_version_index'][msg_key] = len(conv['versions'][msg_key]) - 1

    if conv['title'] == 'New Chat':
        conv['title'] = generate_chat_title(messages)

    conv['last_updated'] = datetime.now().isoformat()
    save_conversations(conversations)

    return jsonify({
        'success': True,
        'new_response': ai_response,
        'conversation_id': conversation_id,
        'messages': messages,
        'versions': conv.get('versions', {}),
        'current_version_index': conv.get('current_version_index', {})
    })

@app.route('/api/switch_version', methods=['POST'])
def switch_version():
    data = request.json
    conversation_id = data.get('conversation_id')
    message_index = data.get('message_index')
    version_index = data.get('version_index')

    if conversation_id not in conversations:
        return jsonify({'error': 'Conversation not found'}), 404

    conv = conversations[conversation_id]
    versions = conv.get('versions', {})
    msg_key = str(message_index)

    if msg_key not in versions or version_index >= len(versions[msg_key]):
        return jsonify({'error': 'Version not found'}), 404

    if 'current_version_index' not in conv:
        conv['current_version_index'] = {}
    
    # Store the previous version index
    previous_version_index = conv['current_version_index'].get(msg_key)
    conv['current_version_index'][msg_key] = version_index

    # CRITICAL FIX: Reconstruct the conversation to show only this version's branch
    # Each version maintains its own independent conversation history
    
    # Get the base conversation up to and including the user message before this assistant response
    base_messages = conv['messages'][:message_index]  # Up to but not including the assistant response
    
    # Add the assistant response for this specific version
    version_response = versions[msg_key][version_index]
    
    # Check if this version has its own branch history stored
    version_branch_key = f"{msg_key}_v{version_index}"
    branch_messages = conv.get('version_branches', {}).get(version_branch_key, [])
    
    # Reconstruct the conversation: base + this version's response + this version's subsequent messages
    new_messages = list(base_messages)
    
    # Add the assistant message for this version
    if message_index < len(conv['messages']):
        assistant_msg = dict(conv['messages'][message_index])  # Copy existing message structure
        assistant_msg['content'] = version_response
        new_messages.append(assistant_msg)
    
    # Add the branch-specific subsequent messages if they exist
    if branch_messages:
        new_messages.extend(branch_messages)
    elif version_index == previous_version_index:
        # If switching to same version, keep current subsequent messages as this version's branch
        subsequent = conv['messages'][message_index + 1:]
        # Store these as this version's branch for future switches
        if 'version_branches' not in conv:
            conv['version_branches'] = {}
        conv['version_branches'][version_branch_key] = list(subsequent)
        new_messages.extend(subsequent)
    
    # Update the conversation with the reconstructed branch
    conv['messages'] = new_messages
    
    # Update branch_root to track which version we're on
    conv['branch_root'] = message_index
    conv['branch_version'] = version_index

    save_conversations(conversations)

    return jsonify({
        'success': True,
        'content': version_response,
        'current_version_index': conv['current_version_index'],
        'messages': conv['messages']
    })

@app.route('/api/rename_chat', methods=['POST'])
def rename_chat():
    data = request.json
    conversation_id = data.get('conversation_id')
    new_title = data.get('new_title')

    if conversation_id not in conversations:
        return jsonify({'error': 'Conversation not found'}), 404

    conversations[conversation_id]['title'] = sanitize_filename(new_title)
    save_conversations(conversations)

    return jsonify({'success': True})

@app.route('/api/export_chat/<conversation_id>', methods=['GET'])
def export_chat(conversation_id):
    if conversation_id not in conversations:
        return jsonify({'error': 'Conversation not found'}), 404

    doc = export_to_word(conversations[conversation_id])

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    # Clean the title to remove any invalid characters for HTTP headers
    title = conversations[conversation_id]['title']
    # Remove newlines, carriage returns, and other control characters
    title = ''.join(char for char in title if char.isprintable() and char not in '\n\r\t')
    # Also remove any other problematic characters
    title = title.replace('/', '-').replace('\\', '-').replace(':', '-')
    
    filename = f"HenAi_Chat_{title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    return send_file(
        file_stream,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/api/get_version', methods=['POST'])
def get_version():
    data = request.json
    conversation_id = data.get('conversation_id')
    message_index = data.get('message_index')
    version_index = data.get('version_index')

    if conversation_id not in conversations:
        return jsonify({'error': 'Conversation not found'}), 404

    conv = conversations[conversation_id]
    versions = conv.get('versions', {})
    msg_key = str(message_index)

    if msg_key not in versions or version_index >= len(versions[msg_key]):
        return jsonify({'error': 'Version not found'}), 404

    return jsonify({
        'success': True,
        'content': versions[msg_key][version_index]
    })

@app.route('/api/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    if not allowed_file(file.filename):
        return jsonify({'error': 'File type not allowed. Please upload text files, code files, or documents.'}), 400

    filename = secure_filename(file.filename)
    file_content = file.read()

    try:
        text_content = extract_text_from_file(file_content, filename)
        
        # Format with header for AI context
        file_extension = filename.rsplit('.', 1)[1].lower() if '.' in filename else 'txt'
        formatted_content = f"\n\n--- FILE: {filename} ({file_extension}) ---\n{text_content}\n--- END FILE: {filename} ---\n\n"

        return jsonify({
            'success': True,
            'filename': filename,
            'content': formatted_content,  # Return formatted content
            'preview': text_content[:500] + ('...' if len(text_content) > 500 else '')
        })
    except Exception as e:
        return jsonify({'error': f'Error processing file: {str(e)}'}), 500

@app.route('/api/add_pending_attachment', methods=['POST'])
def add_pending_attachment():
    """Add a file as a pending attachment to a conversation (not sent to AI yet)"""
    data = request.json
    conversation_id = data.get('conversation_id')
    filename = data.get('filename')
    content = data.get('content')
    
    if not conversation_id or not filename:
        return jsonify({'error': 'Missing conversation_id or filename'}), 400
    
    # Sanitize content to remove invalid Unicode surrogates
    if content:
        content = content.encode('utf-8', errors='replace').decode('utf-8')
    
    # Check if conversation exists
    if conversation_id not in conversations:
        return jsonify({'error': 'Conversation not found'}), 404
    
    conv = conversations[conversation_id]
    
    # Initialize pending_attachments if not exists
    if 'pending_attachments' not in conv:
        conv['pending_attachments'] = []
    
    # Add the file as a pending attachment
    conv['pending_attachments'].append({
        'name': filename,
        'content': content,
        'timestamp': datetime.now().isoformat(),
        'source': 'archive_search'
    })
    
    conv['last_updated'] = datetime.now().isoformat()
    
    # Save to file
    try:
        save_conversations(conversations)
    except Exception as e:
        print(f"Error saving: {e}")
        return jsonify({'error': f'Failed to save attachment: {str(e)}'}), 500
    
    return jsonify({
        'success': True,
        'conversation_id': conversation_id,
        'filename': filename,
        'attachments_count': len(conv['pending_attachments'])
    })

@app.route('/api/attach_to_chat', methods=['POST'])
def attach_to_chat():
    """Attach a file from archive search to a conversation as a pending attachment"""
    data = request.json
    conversation_id = data.get('conversation_id')
    filename = data.get('filename')
    content = data.get('content')
    
    if not conversation_id or not filename:
        return jsonify({'error': 'Missing conversation_id or filename'}), 400
    
    # Sanitize content to remove invalid Unicode surrogates
    if content:
        content = content.encode('utf-8', errors='replace').decode('utf-8')
    
    # Check if conversation exists, if not create it
    if conversation_id not in conversations:
        conversation_id = str(uuid.uuid4())
        conversations[conversation_id] = {
            'id': conversation_id,
            'messages': [],
            'title': 'New Chat',
            'created_at': datetime.now().isoformat(),
            'last_updated': datetime.now().isoformat(),
            'versions': {},
            'current_version_index': {},
            'branch_root': None,
            'pending_attachments': []
        }
    
    conv = conversations[conversation_id]
    
    # Initialize pending_attachments if not exists
    if 'pending_attachments' not in conv:
        conv['pending_attachments'] = []
    
    # Add the file as a pending attachment
    conv['pending_attachments'].append({
        'name': filename,
        'content': content,
        'timestamp': datetime.now().isoformat(),
        'source': 'archive_search'
    })
    
    conv['last_updated'] = datetime.now().isoformat()
    
    # Save to file - handle encoding properly
    try:
        save_conversations(conversations)
    except Exception as e:
        print(f"Error saving: {e}")
        if 'pending_attachments' in conv and conv['pending_attachments']:
            conv['pending_attachments'].pop()
            save_conversations(conversations)
    
    return jsonify({
        'success': True,
        'conversation_id': conversation_id,
        'filename': filename,
        'attachments_count': len(conv['pending_attachments'])
    })

@app.route('/api/conversations', methods=['GET'])
def get_conversations():
    conv_list = []
    for conv_id, conv_data in conversations.items():
        conv_list.append({
            'id': conv_id,
            'title': conv_data.get('title', 'New Chat'),
            'message_count': len(conv_data.get('messages', [])),
            'last_updated': conv_data.get('last_updated', datetime.now().isoformat())
        })
    conv_list.sort(key=lambda x: x['last_updated'], reverse=True)
    return jsonify(conv_list)

@app.route('/api/conversation/<conversation_id>', methods=['GET'])
def get_conversation(conversation_id):
    if conversation_id in conversations:
        conv_data = conversations[conversation_id]
        
        messages = conv_data.get('messages', [])
        for msg in messages:
            if msg.get('media_data'):
                msg['media_data'] = msg['media_data']
        
        return jsonify({
            **conv_data,
            'versions': conv_data.get('versions', {}),
            'current_version_index': conv_data.get('current_version_index', {}),
            'version_branches': conv_data.get('version_branches', {}),
            'branch_root': conv_data.get('branch_root'),
            'branch_version': conv_data.get('branch_version'),
            'pending_attachments': conv_data.get('pending_attachments', [])
        })
    return jsonify({'error': 'Conversation not found'}), 404

@app.route('/api/delete_conversation/<conversation_id>', methods=['DELETE'])
def delete_conversation(conversation_id):
    if conversation_id in conversations:
        del conversations[conversation_id]
        save_conversations(conversations)
        return jsonify({'success': True})
    return jsonify({'error': 'Conversation not found'}), 404

@app.route('/api/remove_pending_attachment', methods=['POST'])
def remove_pending_attachment():
    """Remove a specific pending attachment"""
    data = request.json
    conversation_id = data.get('conversation_id')
    attachment_index = data.get('attachment_index')
    
    if conversation_id not in conversations:
        return jsonify({'error': 'Conversation not found'}), 404
    
    conv = conversations[conversation_id]
    
    if 'pending_attachments' not in conv:
        return jsonify({'error': 'No pending attachments'}), 404
    
    if attachment_index >= len(conv['pending_attachments']):
        return jsonify({'error': 'Attachment not found'}), 404
    
    # Remove the attachment
    conv['pending_attachments'].pop(attachment_index)
    
    # Clean up empty list if needed
    if len(conv['pending_attachments']) == 0:
        conv['pending_attachments'] = []
    
    conv['last_updated'] = datetime.now().isoformat()
    save_conversations(conversations)
    
    return jsonify({'success': True, 'remaining': len(conv['pending_attachments'])})

@app.route('/api/clear_pending_attachments', methods=['POST'])
def clear_pending_attachments():
    """Clear all pending attachments"""
    data = request.json
    conversation_id = data.get('conversation_id')
    
    if conversation_id not in conversations:
        return jsonify({'error': 'Conversation not found'}), 404
    
    conv = conversations[conversation_id]
    conv['pending_attachments'] = []
    conv['last_updated'] = datetime.now().isoformat()
    save_conversations(conversations)
    
    return jsonify({'success': True})

# ============= IDE ENDPOINTS (using workspace.py) =============

from workspace import load_workspace, save_workspace, build_folder_context, get_file_content_from_workspace

@app.route('/api/ide/chat', methods=['POST'])
def ide_chat():
    """Handle AI chat with folder context included"""
    data = request.json
    message = data.get('message', '')
    folder_path = data.get('folder_path', '')
    selected_file = data.get('selected_file', '')
    conversation_history = data.get('conversation_history', [])
    include_full_folder = data.get('include_full_folder', True)
    
    if not message:
        return jsonify({'error': 'No message provided'}), 400
    
    # Build context with folder contents if requested
    context_content = ""
    
    if include_full_folder and folder_path:
        # Get folder tree
        workspace = load_workspace()
        parts = folder_path.split('/') if folder_path else []
        current = workspace
        for part in parts:
            found = False
            for folder in current.get('folders', []):
                if folder.get('name') == part:
                    current = folder
                    found = True
                    break
            if not found:
                break
        
        # Build folder structure description
        context_content = build_folder_context(current, folder_path)
    
    # Add selected file content if specified
    if selected_file:
        file_content = get_file_content_from_workspace(selected_file)
        if file_content:
            context_content += f"\n\n--- SELECTED FILE: {selected_file} ---\n"
            context_content += file_content
            context_content += f"\n--- END FILE: {selected_file} ---\n"
    
    # Prepare messages for AI
    system_prompt = """You are HenAi, an expert coding assistant integrated into the HenIde workspace. 
You have access to the entire folder structure and file contents when the checkbox is enabled.

IMPORTANT RULES:
1. When providing code, output ONLY the code with no explanations before or after
2. Use proper code blocks with language specification
3. Understand the context of the entire folder structure
4. When asked to debug, explain what's wrong and provide the corrected code
5. When asked to generate new files, specify the filename and provide the complete code
6. Keep responses concise but thorough

Remember: Code output should be pure code with no additional text outside code blocks."""
    
    # Build messages with context
    messages = [{"role": "system", "content": system_prompt}]
    
    # Add conversation history
    for msg in conversation_history:
        messages.append({"role": msg.get('role'), "content": msg.get('content')})
    
    # Add folder context if present
    if context_content:
        context_message = f"FOLDER CONTEXT:\n{context_content}\n\nUSER QUERY: {message}"
        messages.append({"role": "user", "content": context_message})
    else:
        messages.append({"role": "user", "content": message})
    
    # Determine if this is a code generation request
    is_code_gen = is_code_generation_request(message)
    
    # Call AI with appropriate service
    ai_response = query_ai_with_fallback(
        messages[-1]['content'] if not context_content else context_message,
        messages[:-1],
        is_code_gen
    )
    
    return jsonify({
        'success': True,
        'response': ai_response
    })

@app.route('/api/ide/create_file', methods=['POST'])
def ide_create_file():
    """Create a new file in the IDE"""
    data = request.json
    folder_path = data.get('folder_path', '')
    file_name = data.get('file_name', '').strip()
    content = data.get('content', '')
    
    if not file_name:
        return jsonify({'error': 'File name required'}), 400
    
    # Validate file name
    if not re.match(r'^[a-zA-Z0-9_\-\.]+$', file_name):
        return jsonify({'error': 'Invalid file name. Use only letters, numbers, underscore, hyphen, and dot.'}), 400
    
    workspace = load_workspace()
    
    # Navigate to the target folder
    if folder_path:
        parts = folder_path.split('/') if folder_path else []
        current = workspace
        for part in parts:
            found = False
            for folder in current.get('folders', []):
                if folder.get('name') == part:
                    current = folder
                    found = True
                    break
            if not found:
                return jsonify({'error': f'Folder "{folder_path}" not found'}), 404
        target_folder = current
    else:
        target_folder = workspace
    
    # Ensure files list exists
    if 'files' not in target_folder:
        target_folder['files'] = []
    
    # Check if file already exists
    for file_item in target_folder['files']:
        if file_item.get('name') == file_name:
            return jsonify({'error': 'File already exists'}), 400
    
    # Create new file (insert at beginning for newest first)
    new_file = {
        'name': file_name,
        'content': content,
        'created_at': datetime.now().isoformat(),
        'updated_at': datetime.now().isoformat()
    }
    target_folder['files'].insert(0, new_file)
    
    save_workspace(workspace)
    
    full_path = folder_path + '/' + file_name if folder_path else file_name
    print(f"✅ IDE: File created successfully - {full_path}")
    
    return jsonify({
        'success': True,
        'file': new_file,
        'full_path': full_path
    })

@app.route('/api/ide/create_folder', methods=['POST'])
def ide_create_folder():
    """Create a new subfolder in the IDE"""
    data = request.json
    parent_path = data.get('parent_path', '')
    folder_name = data.get('folder_name', '').strip()
    
    if not folder_name:
        return jsonify({'error': 'Folder name required'}), 400
    
    # Validate folder name
    if not re.match(r'^[a-zA-Z0-9_\-]+$', folder_name):
        return jsonify({'error': 'Invalid folder name. Use only letters, numbers, underscore, and hyphen.'}), 400
    
    workspace = load_workspace()
    
    # Navigate to parent folder
    if parent_path:
        parts = parent_path.split('/') if parent_path else []
        current = workspace
        for part in parts:
            found = False
            for folder in current.get('folders', []):
                if folder.get('name') == part:
                    current = folder
                    found = True
                    break
            if not found:
                return jsonify({'error': f'Parent folder "{parent_path}" not found'}), 404
        parent = current
    else:
        parent = workspace
    
    # Ensure folders list exists
    if 'folders' not in parent:
        parent['folders'] = []
    
    # Check if folder already exists
    for folder in parent['folders']:
        if folder.get('name') == folder_name:
            return jsonify({'error': 'Folder already exists'}), 400
    
    # Create new folder (insert at beginning for newest first)
    new_folder = {
        'name': folder_name,
        'created_at': datetime.now().isoformat(),
        'folders': [],
        'files': []
    }
    parent['folders'].insert(0, new_folder)
    
    save_workspace(workspace)
    
    new_path = parent_path + '/' + folder_name if parent_path else folder_name
    print(f"✅ IDE: Folder created successfully - {new_path}")
    
    return jsonify({
        'success': True,
        'folder': new_folder,
        'full_path': new_path
    })

@app.route('/api/ide/update_file', methods=['POST'])
def ide_update_file():
    """Update file content from IDE"""
    data = request.json
    file_path = data.get('file_path', '')
    content = data.get('content', '')
    
    if not file_path:
        return jsonify({'error': 'File path required'}), 400
    
    workspace = load_workspace()
    
    # Parse file path
    parts = file_path.split('/')
    file_name = parts[-1]
    folder_parts = parts[:-1] if len(parts) > 1 else []
    
    # Navigate to the folder containing the file
    current = workspace
    for folder_name in folder_parts:
        found = False
        # Search in current's folders
        for folder in current.get('folders', []):
            if folder.get('name') == folder_name:
                current = folder
                found = True
                break
        if not found:
            return jsonify({'error': f'Folder "{folder_name}" not found in path'}), 404
    
    # Find and update the file
    files = current.get('files', [])
    for i, file_item in enumerate(files):
        if file_item.get('name') == file_name:
            file_item['content'] = content
            file_item['updated_at'] = datetime.now().isoformat()
            # Also update the file in the list to ensure reference is kept
            files[i] = file_item
            save_workspace(workspace)
            print(f"✅ IDE: File saved successfully - {file_path}")
            return jsonify({'success': True, 'file': file_item})
    
    return jsonify({'error': f'File "{file_name}" not found in folder'}), 404

@app.route('/api/ide/delete_file', methods=['POST'])
def ide_delete_file():
    """Delete a file from IDE"""
    data = request.json
    file_path = data.get('file_path', '')
    
    if not file_path:
        return jsonify({'error': 'File path required'}), 400
    
    workspace = load_workspace()
    
    # Parse file path
    parts = file_path.split('/')
    file_name = parts[-1]
    folder_parts = parts[:-1] if len(parts) > 1 else []
    
    # Navigate to the folder containing the file
    current = workspace
    for folder_name in folder_parts:
        found = False
        for folder in current.get('folders', []):
            if folder.get('name') == folder_name:
                current = folder
                found = True
                break
        if not found:
            return jsonify({'error': f'Folder not found'}), 404
    
    # Find and delete the file
    files = current.get('files', [])
    for i, file_item in enumerate(files):
        if file_item.get('name') == file_name:
            del files[i]
            save_workspace(workspace)
            return jsonify({'success': True})
    
    return jsonify({'error': 'File not found'}), 404

# ============= MEDIA API ENDPOINTS =============

# Import media handler
from media import media_handler

@app.route('/api/media/search/image', methods=['POST'])
def search_image():
    """Search for a single image with automatic provider fallback"""
    data = request.json
    query = data.get('query', '').strip()
    provider = data.get('provider')  # Optional - if not provided, will use fallback
    
    if not query:
        return jsonify({'error': 'No query provided'}), 400
    
    # If provider is specified, use it directly
    if provider:
        result = media_handler.search_images(query, provider)
    else:
        # Use fallback search that tries random providers (now includes DuckDuckGo and Openverse)
        result = media_handler.search_with_fallback(query, 'image')
    
    return jsonify(result)

@app.route('/api/media/search/video', methods=['POST'])
def search_video():
    """Search for a single video with automatic provider fallback"""
    data = request.json
    query = data.get('query', '').strip()
    provider = data.get('provider')  # Optional - if not provided, will use fallback
    
    if not query:
        return jsonify({'error': 'No query provided'}), 400
    
    # If provider is specified, use it directly
    if provider:
        result = media_handler.search_videos(query, provider)
    else:
        # Use fallback search that tries random providers (now includes DuckDuckGo)
        result = media_handler.search_with_fallback(query, 'video')
    
    return jsonify(result)

@app.route('/api/media/regenerate', methods=['POST'])
def regenerate_media():
    """Regenerate media with automatic fallback to different providers"""
    data = request.json
    query = data.get('query', '').strip()
    media_type = data.get('media_type', 'image')
    current_id = data.get('current_id')
    provider = data.get('provider')
    
    if not query:
        return jsonify({'error': 'No query provided'}), 400
    
    # Use the fallback regeneration method
    result = media_handler.regenerate_with_fallback(query, media_type, current_id, provider)
    
    return jsonify(result)

@app.route('/api/media/analyze/video', methods=['POST'])
def analyze_video():
    """Analyze a video using TwelveLabs"""
    data = request.json
    video_url = data.get('video_url', '')
    video_name = data.get('video_name', 'video.mp4')
    
    if not video_url:
        return jsonify({'error': 'No video URL provided'}), 400
    
    result = media_handler.analyze_video(video_url, video_name)
    return jsonify(result)

@app.route('/api/media/analyze/image', methods=['POST'])
def analyze_image():
    """Analyze an image using Vision Model (BLIP/Florence) + OCR as supplement"""
    data = request.json
    image_url = data.get('image_url', '')
    image_name = data.get('image_name', 'image.jpg')
    
    if not image_url:
        return jsonify({'error': 'No image URL provided'}), 400
    
    temp_file_path = None
    try:
        # Download the image
        print(f"📥 Downloading image from: {image_url}")
        response = requests.get(image_url, timeout=30, stream=True)
        response.raise_for_status()
        
        # Get image content
        image_content = response.content
        
        # Save to temporary file for processing (optional, for OCR)
        import tempfile
        import re
        
        # Determine file extension
        ext = image_name.split('.')[-1].lower() if '.' in image_name else 'jpg'
        if ext not in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'webp']:
            ext = 'jpg'
        
        with tempfile.NamedTemporaryFile(suffix=f'.{ext}', delete=False) as tmp:
            tmp.write(image_content)
            tmp.flush()
            temp_file_path = tmp.name
        
        print(f"📁 Image saved to temp file: {temp_file_path}")
        
        # ============= STEP 1: VISION MODEL (Primary) =============
        from vision import get_vision_model
        vision_model = get_vision_model()
        
        print("🔍 Analyzing image content with Vision Model...")
        vision_analysis = vision_model.analyze_image(image_content)
        
        # ============= STEP 2: OCR TEXT (Secondary, for text-heavy images) =============
        ocr_text = ""
        try:
            # Only run OCR if vision model gave a short result or failed
            # This saves time on non-text images
            if not vision_analysis or len(vision_analysis) < 20:
                print("📝 Running OCR for text extraction...")
                processed_output = binary_processor.process_image(image_content, image_name)
                
                # Extract just the OCR text, not the full metadata
                if "EXTRACTED TEXT:" in processed_output:
                    ocr_match = re.search(r'EXTRACTED TEXT:\n(.*?)(?=\n--- END IMAGE ANALYSIS|--- END FILE)', processed_output, re.DOTALL)
                    if ocr_match:
                        ocr_text = ocr_match.group(1).strip()
                        print(f"✅ OCR extracted {len(ocr_text)} characters")
                elif "No text detected" not in processed_output:
                    # Try to extract any meaningful text
                    lines = processed_output.split('\n')
                    for i, line in enumerate(lines):
                        if '📝 EXTRACTED TEXT:' in line and i + 1 < len(lines):
                            ocr_text = lines[i + 1].strip()
                            break
            else:
                print("✅ Vision model provided good analysis, skipping OCR")
        except Exception as e:
            print(f"⚠️ OCR extraction skipped: {e}")
        
        # ============= STEP 3: Build Clean Analysis =============
        clean_analysis = ""
        
        # Primary: Use vision model result
        if vision_analysis and len(vision_analysis) > 10:
            clean_analysis = vision_analysis
            print(f"✅ Using Vision Model analysis ({len(clean_analysis)} chars)")
            
            # Supplement with OCR if there's meaningful text (like labels, signs)
            if ocr_text and len(ocr_text.strip()) > 5 and not ocr_text.startswith("[OCR"):
                # Check if OCR text is substantial (not just a single character)
                if len(ocr_text) > 10:
                    clean_analysis += f"\n\nThe image also contains visible text: \"{ocr_text[:200]}\""
                else:
                    clean_analysis += f" The text \"{ocr_text.strip()}\" appears in the image."
        
        # Fallback: Use OCR only
        elif ocr_text and len(ocr_text.strip()) > 10 and not ocr_text.startswith("[OCR"):
            clean_analysis = f"The image contains readable text:\n\n{ocr_text[:500]}"
            print("⚠️ Using OCR fallback")
        
        # Final fallback: Use filename hint
        else:
            name_without_ext = re.sub(r'\.[^.]+$', '', image_name)
            clean_name = re.sub(r'[_\-\.]', ' ', name_without_ext)
            clean_name = re.sub(r'\d+', '', clean_name).strip()
            if clean_name and len(clean_name) > 3:
                clean_analysis = f"This image appears to show {clean_name}."
            else:
                clean_analysis = "The image has been processed, but no readable text or recognizable content was detected."
            print("⚠️ Using filename fallback")
        
        # ============= STEP 4: Final Cleanup - Remove Any Metadata =============
        # Remove common metadata patterns
        metadata_patterns = [
            r'Photographer:?\s*\S+',  # Photographer name
            r'Photo by\s+\S+',  # Photo by [name]
            r'Credit:?\s*\S+',  # Credit:
            r'Source:?\s*\S+',  # Source:
            r'Copyright\s+[©]?\s*\S+',  # Copyright notices
            r'©\s*\d{4}\s*\S+',  # © year name
            r'\[.*?\]',  # Any bracketed text that might be metadata
            r'\(.*?(credit|courtesy|source).*?\)',  # Parenthetical credits
            r'Image courtesy of\s+\S+',  # Courtesy notices
            r'Sourced from\s+\S+',  # Sourced from
            r'^\w+:\s*$',  # Empty labels at line start
            r'\*\*Image Analysis:.*?\*\*',  # Markdown headers
            r'\*\*AI Analysis:\*\*',  # AI Analysis header
            r'\*\*OCR Extracted Text Found:\*\*',  # OCR header
            r'\*\*Note:\*\*',  # Note header
            r'\*\*Image Details:\*\*',  # Image Details header
            r'\*\*Final Analysis:\*\*',  # Final Analysis header
            r'---.*?---',  # Horizontal rules
        ]
        
        for pattern in metadata_patterns:
            clean_analysis = re.sub(pattern, '', clean_analysis, flags=re.IGNORECASE | re.MULTILINE)
        
        # Remove markdown formatting
        clean_analysis = re.sub(r'\*\*([^*]+)\*\*', r'\1', clean_analysis)  # Bold
        clean_analysis = re.sub(r'`([^`]+)`', r'\1', clean_analysis)  # Code
        clean_analysis = re.sub(r'#+\s*', '', clean_analysis)  # Headers
        
        # Clean up extra whitespace
        clean_analysis = re.sub(r'\s+', ' ', clean_analysis)
        clean_analysis = re.sub(r'\n{3,}', '\n\n', clean_analysis)
        
        # Ensure first letter is capitalized
        if clean_analysis and len(clean_analysis) > 0:
            clean_analysis = clean_analysis[0].upper() + clean_analysis[1:] if len(clean_analysis) > 1 else clean_analysis.upper()
        
        # Remove any trailing punctuation that looks like incomplete sentences
        clean_analysis = re.sub(r'\s*[|;:]\s*$', '', clean_analysis)
        
        # Final trim
        clean_analysis = clean_analysis.strip()
        
        # If analysis is empty after cleaning, provide a generic response
        if not clean_analysis or len(clean_analysis) < 5:
            clean_analysis = "The image has been analyzed, but no specific content could be identified."
        
        response_data = {
            'success': True,
            'image_url': image_url,
            'image_name': image_name,
            'analysis': clean_analysis
        }
        
        print(f"✅ Analysis complete: {clean_analysis[:100]}...")
        return jsonify(response_data)
        
    except requests.exceptions.RequestException as e:
        print(f"❌ Failed to download image: {e}")
        return jsonify({
            'success': False,
            'error': f'Failed to download image: {str(e)}'
        }), 500
    except Exception as e:
        print(f"❌ Error analyzing image: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': f'Error analyzing image: {str(e)}'
        }), 500
    finally:
        # Clean up temp file
        if temp_file_path:
            try:
                import os
                os.unlink(temp_file_path)
            except:
                pass

@app.route('/api/save_media_message', methods=['POST'])
def save_media_message():
    data = request.json
    conversation_id = data.get('conversation_id')
    user_message = data.get('user_message', '')
    assistant_response = data.get('assistant_response', '')
    media_data = data.get('media_data')
    is_error = data.get('is_error', False)
    
    if not conversation_id or conversation_id not in conversations:
        conversation_id = str(uuid.uuid4())
        conversations[conversation_id] = {
            'id': conversation_id,
            'messages': [],
            'title': 'New Chat',
            'created_at': datetime.now().isoformat(),
            'last_updated': datetime.now().isoformat(),
            'versions': {},
            'current_version_index': {},
            'branch_root': None,
            'pending_attachments': []
        }
    
    conv = conversations[conversation_id]
    
    user_msg = {
        'role': 'user',
        'content': user_message,
        'timestamp': datetime.now().isoformat(),
        'attachments': None
    }
    conv['messages'].append(user_msg)
    
    assistant_msg = {
        'role': 'assistant',
        'content': assistant_response,
        'timestamp': datetime.now().isoformat(),
        'media_data': media_data,
        'is_media_result': True,
        'is_error': is_error
    }
    conv['messages'].append(assistant_msg)
    
    if conv['title'] == 'New Chat' and user_message:
        raw_title = generate_chat_title(conv['messages'])
        conv['title'] = sanitize_filename(raw_title)
    
    conv['last_updated'] = datetime.now().isoformat()
    save_conversations(conversations)
    
    return jsonify({
        'success': True,
        'conversation_id': conversation_id,
        'title': conv['title']
    })

# ============= DOCS API ENDPOINTS (using mydocs.py) =============

# Initialize document processor and creator
doc_processor = DocumentProcessor()
doc_creator = DocumentCreator()

@app.route('/api/docs/generate', methods=['POST'])
def docs_generate():
    """Generate document content using AI"""
    data = request.json
    prompt = data.get('prompt', '')
    doc_type = data.get('doc_type', 'word')
    template_id = data.get('template_id')
    
    if not prompt:
        return jsonify({'error': 'No prompt provided'}), 400
    
    # Enhance prompt with template-specific instructions using DocumentCreator
    enhanced_prompt = doc_creator._apply_template_formatting(prompt, doc_type, template_id) if template_id else prompt
    
    try:
        # Use OpenRouter for content generation
        from models import query_openrouter
        response = query_openrouter(enhanced_prompt, context=None, is_code_generation=False)
        
        if response:
            return jsonify({
                'success': True,
                'content': response
            })
        else:
            return jsonify({'error': 'AI generation failed'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/docs/create', methods=['POST'])
def docs_create():
    """Create a document from generated content using DocumentCreator"""
    data = request.json
    doc_type = data.get('doc_type', 'word')
    content = data.get('content', '')
    filename = data.get('filename', 'document')
    template_id = data.get('template_id')
    
    if not content:
        return jsonify({'error': 'No content provided'}), 400
    
    try:
        # Use DocumentCreator to create the document
        output_path = doc_creator.create_document(content, doc_type, filename, template_id)
        
        if output_path and output_path.exists():
            # Return download URL
            download_url = f'/api/docs/download/{output_path.name}'
            return jsonify({
                'success': True,
                'filename': output_path.name,
                'download_url': download_url
            })
        else:
            return jsonify({'error': 'Failed to create document'}), 500
            
    except Exception as e:
        print(f"Error creating document: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/docs/download/<filename>', methods=['GET'])
def docs_download(filename):
    """Download a generated document"""
    file_path = doc_creator.output_dir / filename
    if file_path.exists():
        return send_file(
            file_path,
            as_attachment=True,
            download_name=filename,
            mimetype='application/octet-stream'
        )
    return jsonify({'error': 'File not found'}), 404

@app.route('/api/docs/extract', methods=['POST'])
def docs_extract():
    """Extract content from a document"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    extract_type = request.form.get('extract_type', 'text')
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    filename = secure_filename(file.filename)
    file_content = file.read()
    
    try:
        # Create a temporary file
        temp_path = doc_creator.temp_dir / filename
        with open(temp_path, 'wb') as f:
            f.write(file_content)
        
        # Extract based on file type and extract_type
        extracted_content = ""
        
        suffix = os.path.splitext(filename)[1].lower()
        
        if suffix in ['.docx', '.doc']:
            if extract_type == 'text':
                result = doc_processor.read_word_document(str(temp_path))
                extracted_content = '\n'.join(result['paragraphs'])
            elif extract_type == 'metadata':
                result = doc_processor.read_word_document(str(temp_path))
                extracted_content = f"Paragraphs: {result['metadata']['paragraph_count']}\nTables: {result['metadata']['table_count']}"
            else:
                extracted_content = doc_processor.extract_text_from_file(str(temp_path))
                
        elif suffix == '.pdf':
            if extract_type == 'text':
                result = doc_processor.read_pdf(str(temp_path), extract_tables=False)
                extracted_content = '\n'.join([page['text'] for page in result['pages']])
            elif extract_type == 'metadata':
                result = doc_processor.read_pdf(str(temp_path), extract_tables=False)
                extracted_content = f"Pages: {result['page_count']}\nMetadata: {result.get('metadata', {})}"
            else:
                extracted_content = doc_processor.extract_text_from_file(str(temp_path))
                
        elif suffix in ['.pptx', '.ppt']:
            result = doc_processor.read_presentation(str(temp_path))
            if extract_type == 'slides':
                for i, slide in enumerate(result['slides']):
                    extracted_content += f"\n--- Slide {i+1} ---\n{slide['text_content']}\n"
            elif extract_type == 'notes':
                for i, slide in enumerate(result['slides']):
                    if slide.get('notes'):
                        extracted_content += f"\n--- Slide {i+1} Notes ---\n{slide['notes']}\n"
            else:
                for slide in result['slides']:
                    extracted_content += slide['text_content'] + '\n'
                    
        elif suffix in ['.xlsx', '.xls', '.csv']:
            import pandas as pd
            if suffix == '.csv':
                df = pd.read_csv(str(temp_path))
            else:
                df = pd.read_excel(str(temp_path))
            extracted_content = df.to_string()
            
        elif suffix in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
            extracted_content = doc_processor.extract_text_from_image(str(temp_path))
            
        elif suffix in ['.mp4', '.avi', '.mov', '.mkv']:
            info = doc_processor.extract_video_info(str(temp_path))
            extracted_content = json.dumps(info, indent=2)
            
        else:
            extracted_content = doc_processor.extract_text_from_file(str(temp_path))
        
        return jsonify({
            'success': True,
            'content': extracted_content[:50000],  # Limit size
            'filename': filename
        })
        
    except Exception as e:
        print(f"Error extracting: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500
    finally:
        # Cleanup temp file
        if os.path.exists(temp_path):
            os.unlink(temp_path)

@app.route('/api/docs/merge', methods=['POST'])
def docs_merge():
    """Merge multiple PDF files"""
    files = request.files.getlist('files')
    output_name = request.form.get('output_name', 'merged_document')
    
    if len(files) < 2:
        return jsonify({'error': 'Need at least 2 files to merge'}), 400
    
    try:
        # Save all files temporarily
        temp_files = []
        for file in files:
            if file.filename and file.filename.lower().endswith('.pdf'):
                temp_path = doc_creator.temp_dir / secure_filename(file.filename)
                file.save(str(temp_path))
                temp_files.append(str(temp_path))
        
        # Merge PDFs
        output_filename = f"{output_name}.pdf"
        result_path = doc_processor.merge_pdfs(temp_files, output_filename)
        
        # Cleanup temp files
        for temp_file in temp_files:
            if os.path.exists(temp_file):
                os.unlink(temp_file)
        
        download_url = f'/api/docs/download/{output_filename}'
        return jsonify({
            'success': True,
            'filename': output_filename,
            'download_url': download_url
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/docs/split', methods=['POST'])
def docs_split():
    """Split a PDF file"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    pages_per_file = int(request.form.get('pages_per_file', 1))
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not file.filename.lower().endswith('.pdf'):
        return jsonify({'error': 'File must be PDF'}), 400
    
    try:
        # Save file temporarily
        temp_path = doc_creator.temp_dir / secure_filename(file.filename)
        file.save(str(temp_path))
        
        # Split PDF
        split_dir = doc_creator.temp_dir / 'split_output'
        split_files = doc_processor.split_pdf(str(temp_path), str(split_dir), pages_per_file)
        
        # Prepare response
        result_files = []
        for split_file in split_files:
            filename = os.path.basename(split_file)
            result_files.append({
                'name': filename,
                'url': f'/api/docs/download/{filename}'
            })
            # Move to output directory
            import shutil
            shutil.move(split_file, str(doc_creator.output_dir / filename))
        
        # Cleanup
        os.unlink(temp_path)
        
        return jsonify({
            'success': True,
            'files_count': len(result_files),
            'files': result_files
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/docs/convert', methods=['POST'])
def docs_convert():
    """Convert document between formats"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    to_format = request.form.get('to_format', 'txt')
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    filename = secure_filename(file.filename)
    file_content = file.read()
    
    try:
        # Save file temporarily
        temp_path = doc_creator.temp_dir / filename
        with open(temp_path, 'wb') as f:
            f.write(file_content)
        
        # Determine output extension
        ext_map = {
            'txt': '.txt', 'docx': '.docx', 'pdf': '.pdf',
            'xlsx': '.xlsx', 'csv': '.csv', 'html': '.html',
            'jpg': '.jpg', 'png': '.png', 'mp3': '.mp3'
        }
        output_ext = ext_map.get(to_format, '.txt')
        output_filename = f"{os.path.splitext(filename)[0]}{output_ext}"
        
        # Convert
        output_path = doc_processor.convert_document(str(temp_path), output_filename, output_format=to_format)
        
        # Cleanup
        os.unlink(temp_path)
        
        download_url = f'/api/docs/download/{os.path.basename(output_path)}'
        return jsonify({
            'success': True,
            'filename': os.path.basename(output_path),
            'format': to_format,
            'download_url': download_url
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/docs/resize', methods=['POST'])
def docs_resize():
    """Resize an image"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    width = request.form.get('width')
    height = request.form.get('height')
    maintain_aspect = request.form.get('maintain_aspect', 'true').lower() == 'true'
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    filename = secure_filename(file.filename)
    file_content = file.read()
    
    try:
        # Save file temporarily
        temp_path = doc_creator.temp_dir / filename
        with open(temp_path, 'wb') as f:
            f.write(file_content)
        
        # Resize
        w = int(width) if width else None
        h = int(height) if height else None
        output_filename = f"resized_{filename}"
        
        output_path = doc_processor.resize_image(
            str(temp_path), output_filename,
            width=w, height=h, maintain_aspect=maintain_aspect
        )
        
        # Get new dimensions
        from PIL import Image
        img = Image.open(output_path)
        new_width, new_height = img.size
        
        # Cleanup
        os.unlink(temp_path)
        
        download_url = f'/api/docs/download/{output_filename}'
        return jsonify({
            'success': True,
            'filename': output_filename,
            'width': new_width,
            'height': new_height,
            'download_url': download_url
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# Register workspace routes
register_workspace_routes(app)

@app.route('/api/generated_image/<filename>')
def serve_generated_image(filename):
    """Serve generated images from the output directory"""
    from flask import send_from_directory
    import os
    
    # Security: ensure filename doesn't contain path traversal
    if '..' in filename or filename.startswith('/'):
        return jsonify({'error': 'Invalid filename'}), 400
    
    # Serve from the generated_images directory
    image_dir = os.path.join(os.path.dirname(__file__), 'generated_images')
    return send_from_directory(image_dir, filename)

if __name__ == '__main__':
    print("\n" + "="*60)
    print("🐔 HenAi Server Started!")
    print("="*60)
    print("📍 Visit: http://localhost:5001")
    print("🤖 AI Mode: Hybrid - Pollinations.ai for conversations, OpenRouter for code")
    print("   📝 Conversations/Explanations: Pollinations.ai (fast, prioritized)")
    print("   💻 Code Generation: OpenRouter only (with multiple free models)")
    print("🖼️ Media Search: Images & Videos from Pixabay and Pexels")
    print("🎬 Video Analysis: TwelveLabs API")
    print("💡 Commands: /search, /extract, /code, /image, /help")
    print("📋 Features: Copy, Edit, Regenerate, Version Toggle, Rename, Export")
    print("📎 File Support: Text, Code, Documents (PDF, DOCX, etc.)")
    print("📄 Document Creation: Word, Excel, PowerPoint, PDF, Images")
    print("="*60)
    print("\n🔄 KEEP-ALIVE SYSTEM:")
    print(f"   • Wakeup pings: {'ENABLED' if WAKEUP_ENABLED else 'DISABLED'}")
    print(f"   • Ping interval: {PING_INTERVAL} seconds")
    print(f"   • Health endpoint: /api/health")
    print(f"   • Ping endpoint: /api/ping")
    print("="*60 + "\n")
    
    port = int(os.environ.get("PORT", 5001))
    app.run(debug=False, host='0.0.0.0', port=port)
